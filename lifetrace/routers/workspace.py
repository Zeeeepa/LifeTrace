"""工作区文件管理相关路由"""

import os
from pathlib import Path

from fastapi import APIRouter, File, HTTPException, Query, UploadFile
from fastapi.responses import StreamingResponse

from lifetrace.routers import dependencies as deps
from lifetrace.schemas.workspace import (
    CreateFileRequest,
    CreateFileResponse,
    CreateFolderRequest,
    CreateFolderResponse,
    DeleteFileRequest,
    DeleteFileResponse,
    DocumentAction,
    DocumentAIRequest,
    DocumentAIResponse,
    FileContentResponse,
    FileNode,
    RenameFileRequest,
    RenameFileResponse,
    SaveFileRequest,
    SaveFileResponse,
    UploadFileResponse,
    WorkspaceFilesResponse,
)
from lifetrace.util.logging_config import get_logger
from lifetrace.util.prompt_loader import get_prompt
from lifetrace.util.token_usage_logger import log_token_usage

logger = get_logger()

router = APIRouter(prefix="/api/workspace", tags=["workspace"])


def build_file_tree(root_path: str, parent_id: str = "") -> list[FileNode]:
    """递归构建文件树

    Args:
        root_path: 根目录路径
        parent_id: 父节点ID

    Returns:
        文件节点列表
    """
    nodes = []
    root = Path(root_path)

    if not root.exists():
        return nodes

    # 获取所有子项并排序（文件夹优先，然后按名称排序）
    items = sorted(root.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))

    for item in items:
        # 跳过隐藏文件和目录
        if item.name.startswith("."):
            continue

        # 生成节点ID（使用相对路径）
        relative_path = item.relative_to(Path(deps.config.workspace_dir))
        node_id = str(relative_path).replace(os.sep, "/")

        if item.is_dir():
            # 递归处理子目录
            children = build_file_tree(str(item), node_id)
            node = FileNode(
                id=node_id,
                name=item.name,
                type="folder",
                children=children,
                parent_id=parent_id if parent_id else None,
            )
        else:
            node = FileNode(
                id=node_id,
                name=item.name,
                type="file",
                parent_id=parent_id if parent_id else None,
            )

        nodes.append(node)

    return nodes


def count_files(nodes: list[FileNode]) -> int:
    """统计文件总数（不包括文件夹）"""
    count = 0
    for node in nodes:
        if node.type == "file":
            count += 1
        elif node.children:
            count += count_files(node.children)
    return count


def extract_text_from_docx(file_path: Path) -> str:
    """从 docx 文件中提取文本内容

    Args:
        file_path: docx 文件路径

    Returns:
        提取的文本内容
    """
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    except ImportError:
        logger.error("python-docx 未安装，无法解析 docx 文件")
        raise
    except Exception as e:
        logger.error(f"解析 docx 文件失败: {e}")
        raise


def extract_text_from_doc(file_path: Path) -> str:
    """从 doc 文件中提取文本内容（使用 antiword 或转换方法）

    Args:
        file_path: doc 文件路径

    Returns:
        提取的文本内容

    Note:
        老版本 .doc 文件较难处理，这里尝试用 python-docx 读取，
        如果失败则返回提示信息
    """
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception:
        # 老版本 .doc 文件 python-docx 无法直接读取
        return f"[无法解析旧版 .doc 文件: {file_path.name}，请转换为 .docx 格式]"


ALLOWED_EXTENSIONS = {".txt", ".md", ".doc", ".docx"}


def _get_unique_file_path(target_dir: Path, filename: str) -> Path:
    """获取唯一的文件路径，如果文件已存在则添加数字后缀"""
    target_file = target_dir / filename
    if not target_file.exists():
        return target_file

    base_name = target_file.stem
    ext = target_file.suffix
    counter = 1
    while target_file.exists():
        target_file = target_dir / f"{base_name}_{counter}{ext}"
        counter += 1
    return target_file


def _decode_text_content(file_content: bytes) -> str:
    """尝试多种编码解码文本内容"""
    encodings = ["utf-8", "gbk", "latin-1"]
    for encoding in encodings:
        try:
            return file_content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_content.decode("latin-1")


def _process_text_file(target_file: Path, file_content: bytes) -> str:
    """处理文本文件（txt, md）"""
    content = _decode_text_content(file_content)
    target_file.write_text(content, encoding="utf-8")
    return content


def _process_doc_file(target_file: Path, file_content: bytes, is_docx: bool) -> tuple[Path, str]:
    """处理 Word 文档文件（doc, docx），返回 (最终文件路径, 内容)"""
    target_file.write_bytes(file_content)
    content = extract_text_from_docx(target_file) if is_docx else extract_text_from_doc(target_file)
    md_file = target_file.with_suffix(".md")
    md_file.write_text(content, encoding="utf-8")
    target_file.unlink()
    return md_file, content


def _validate_upload_target(folder: str, workspace_dir: str) -> tuple[Path | None, str | None]:
    """验证并构建上传目标目录，返回 (目标目录, 错误信息)"""
    target_dir = Path(workspace_dir) / folder if folder else Path(workspace_dir)
    target_dir.mkdir(parents=True, exist_ok=True)

    try:
        target_dir.resolve().relative_to(Path(workspace_dir).resolve())
    except ValueError:
        return None, "禁止访问工作区外的目录"

    return target_dir, None


@router.post("/upload", response_model=UploadFileResponse)
async def upload_file(
    file: UploadFile = File(...),
    folder: str = Query("", description="目标文件夹路径（相对于 workspace）"),
):
    """上传文件到工作区

    支持的文件格式：txt, md, doc, docx

    Args:
        file: 上传的文件
        folder: 目标文件夹路径（可选）
    """
    try:
        if not file.filename:
            return UploadFileResponse(success=False, error="文件名不能为空")

        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            return UploadFileResponse(
                success=False,
                error=f"不支持的文件格式: {file_ext}，支持的格式: {', '.join(ALLOWED_EXTENSIONS)}",
            )

        workspace_dir = deps.config.workspace_dir
        target_dir, error = _validate_upload_target(folder, workspace_dir)
        if error:
            return UploadFileResponse(success=False, error=error)

        target_file = _get_unique_file_path(target_dir, file.filename)
        file_content = await file.read()

        # 根据文件类型处理内容
        if file_ext in {".txt", ".md"}:
            content = _process_text_file(target_file, file_content)
        elif file_ext in {".docx", ".doc"}:
            target_file, content = _process_doc_file(target_file, file_content, file_ext == ".docx")
        else:
            content = ""

        relative_path = target_file.relative_to(Path(workspace_dir))
        file_id = str(relative_path).replace(os.sep, "/")
        logger.info(f"文件上传成功: {file_id}")

        return UploadFileResponse(
            success=True,
            file_id=file_id,
            file_name=target_file.name,
            content=content,
        )

    except Exception as e:
        logger.error(f"文件上传失败: {e}")
        return UploadFileResponse(success=False, error=str(e))


@router.post("/create", response_model=CreateFileResponse)
async def create_file(request: CreateFileRequest):
    """创建新文件

    Args:
        request: 创建文件请求，包含文件名和可选的文件夹路径
    """
    try:
        if not request.file_name:
            return CreateFileResponse(success=False, error="文件名不能为空")

        workspace_dir = deps.config.workspace_dir

        # 构建目标文件路径
        if request.folder:
            target_dir = Path(workspace_dir) / request.folder
        else:
            target_dir = Path(workspace_dir)

        # 确保目标目录存在
        target_dir.mkdir(parents=True, exist_ok=True)

        # 安全检查：确保目标目录在 workspace 内
        try:
            target_dir.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            return CreateFileResponse(success=False, error="禁止访问工作区外的目录")

        # 构建文件路径
        target_file = target_dir / request.file_name

        # 如果文件已存在，添加数字后缀
        if target_file.exists():
            base_name = target_file.stem
            ext = target_file.suffix
            counter = 1
            while target_file.exists():
                target_file = target_dir / f"{base_name}_{counter}{ext}"
                counter += 1

        # 创建文件
        target_file.write_text(request.content, encoding="utf-8")

        # 计算文件 ID（相对路径）
        relative_path = target_file.relative_to(Path(workspace_dir))
        file_id = str(relative_path).replace(os.sep, "/")

        logger.info(f"文件创建成功: {file_id}")

        return CreateFileResponse(
            success=True,
            file_id=file_id,
            file_name=target_file.name,
        )

    except Exception as e:
        logger.error(f"文件创建失败: {e}")
        return CreateFileResponse(success=False, error=str(e))


@router.post("/create-folder", response_model=CreateFolderResponse)
async def create_folder(request: CreateFolderRequest):
    """创建新文件夹

    Args:
        request: 创建文件夹请求，包含文件夹名和可选的父文件夹路径
    """
    try:
        if not request.folder_name:
            return CreateFolderResponse(success=False, error="文件夹名不能为空")

        workspace_dir = deps.config.workspace_dir

        # 构建目标文件夹路径
        if request.parent_folder:
            target_dir = Path(workspace_dir) / request.parent_folder
        else:
            target_dir = Path(workspace_dir)

        # 确保父目录存在
        target_dir.mkdir(parents=True, exist_ok=True)

        # 安全检查：确保目标目录在 workspace 内
        try:
            target_dir.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            return CreateFolderResponse(success=False, error="禁止访问工作区外的目录")

        # 构建文件夹路径
        target_folder = target_dir / request.folder_name

        # 如果文件夹已存在，添加数字后缀
        if target_folder.exists():
            base_name = request.folder_name
            counter = 1
            while target_folder.exists():
                target_folder = target_dir / f"{base_name}_{counter}"
                counter += 1

        # 创建文件夹
        target_folder.mkdir(parents=True, exist_ok=True)

        # 计算文件夹 ID（相对路径）
        relative_path = target_folder.relative_to(Path(workspace_dir))
        folder_id = str(relative_path).replace(os.sep, "/")

        logger.info(f"文件夹创建成功: {folder_id}")

        return CreateFolderResponse(
            success=True,
            folder_id=folder_id,
            folder_name=target_folder.name,
        )

    except Exception as e:
        logger.error(f"文件夹创建失败: {e}")
        return CreateFolderResponse(success=False, error=str(e))


@router.get("/files", response_model=WorkspaceFilesResponse)
async def get_workspace_files():
    """获取工作区文件列表"""
    try:
        workspace_dir = deps.config.workspace_dir

        # 确保目录存在
        if not os.path.exists(workspace_dir):
            logger.warning(f"工作区目录不存在: {workspace_dir}")
            return WorkspaceFilesResponse(files=[], total=0)

        # 构建文件树
        files = build_file_tree(workspace_dir)
        total = count_files(files)

        logger.debug(f"获取工作区文件列表成功，共 {total} 个文件")
        return WorkspaceFilesResponse(files=files, total=total)

    except Exception as e:
        logger.error(f"获取工作区文件列表失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


@router.get("/file", response_model=FileContentResponse)
async def get_file_content(file_id: str = Query(..., description="文件ID（相对路径）")):
    """获取文件内容

    Args:
        file_id: 文件ID，即相对于workspace目录的路径
    """
    try:
        workspace_dir = deps.config.workspace_dir

        # 构建文件完整路径
        file_path = Path(workspace_dir) / file_id

        # 安全检查：确保路径在workspace目录内
        try:
            file_path.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            raise HTTPException(status_code=403, detail="禁止访问工作区外的文件") from None

        # 检查文件是否存在
        if not file_path.exists():
            raise HTTPException(status_code=404, detail=f"文件不存在: {file_id}")

        # 检查是否是文件
        if not file_path.is_file():
            raise HTTPException(status_code=400, detail=f"不是文件: {file_id}")

        # 读取文件内容
        try:
            content = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            # 如果无法以UTF-8读取，尝试以二进制读取并返回base64
            import base64

            content = base64.b64encode(file_path.read_bytes()).decode("ascii")

        return FileContentResponse(
            id=file_id,
            name=file_path.name,
            content=content,
            type="file",
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取文件内容失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


def _validate_rename_request(
    request: RenameFileRequest, old_path: Path, workspace_dir: str
) -> str | None:
    """验证重命名请求，返回错误信息或 None"""
    # 安全检查：确保路径在 workspace 目录内
    try:
        old_path.resolve().relative_to(Path(workspace_dir).resolve())
    except ValueError:
        return "禁止访问工作区外的文件"

    # 检查文件是否存在
    if not old_path.exists():
        return f"文件不存在: {request.file_id}"

    # 验证新文件名
    new_name = request.new_name.strip()
    if not new_name:
        return "文件名不能为空"

    # 检查新文件名是否包含非法字符
    invalid_chars = ["/", "\\", ":", "*", "?", '"', "<", ">", "|"]
    if any(char in new_name for char in invalid_chars):
        return f"文件名包含非法字符: {invalid_chars}"

    # 构建新路径并检查是否已存在
    new_path = old_path.parent / new_name
    if new_path.exists() and new_path != old_path:
        return f"文件已存在: {new_name}"

    return None


@router.post("/rename", response_model=RenameFileResponse)
async def rename_file(request: RenameFileRequest):
    """重命名文件或文件夹

    Args:
        request: 重命名请求，包含文件 ID 和新名称
    """
    try:
        workspace_dir = deps.config.workspace_dir
        old_path = Path(workspace_dir) / request.file_id

        # 验证请求
        error = _validate_rename_request(request, old_path, workspace_dir)
        if error:
            return RenameFileResponse(
                success=False,
                old_id=request.file_id,
                new_id=request.file_id,
                new_name=request.new_name,
                error=error,
            )

        # 执行重命名
        new_name = request.new_name.strip()
        new_path = old_path.parent / new_name
        old_path.rename(new_path)

        # 计算新的文件 ID
        new_relative_path = new_path.relative_to(Path(workspace_dir))
        new_id = str(new_relative_path).replace(os.sep, "/")

        logger.info(f"文件重命名成功: {request.file_id} -> {new_id}")

        return RenameFileResponse(
            success=True,
            old_id=request.file_id,
            new_id=new_id,
            new_name=new_name,
        )

    except Exception as e:
        logger.error(f"重命名文件失败: {e}")
        return RenameFileResponse(
            success=False,
            old_id=request.file_id,
            new_id=request.file_id,
            new_name=request.new_name,
            error=str(e),
        )


@router.post("/save", response_model=SaveFileResponse)
async def save_file(request: SaveFileRequest):
    """保存文件内容

    Args:
        request: 保存文件请求，包含文件ID和内容
    """
    try:
        workspace_dir = deps.config.workspace_dir
        file_path = Path(workspace_dir) / request.file_id

        # 检查文件是否存在
        if not file_path.exists():
            return SaveFileResponse(
                success=False,
                file_id=request.file_id,
                error="文件不存在",
            )

        # 检查是否为文件（不是目录）
        if not file_path.is_file():
            return SaveFileResponse(
                success=False,
                file_id=request.file_id,
                error="路径不是文件",
            )

        # 写入文件内容
        file_path.write_text(request.content, encoding="utf-8")

        # 获取更新时间
        from datetime import datetime

        updated_at = datetime.now().isoformat()

        logger.info(f"保存文件成功: {request.file_id}")
        return SaveFileResponse(
            success=True,
            file_id=request.file_id,
            updated_at=updated_at,
        )

    except Exception as e:
        logger.error(f"保存文件失败: {e}")
        return SaveFileResponse(
            success=False,
            file_id=request.file_id,
            error=str(e),
        )


@router.post("/delete", response_model=DeleteFileResponse)
async def delete_file(request: DeleteFileRequest):
    """删除文件或文件夹

    Args:
        request: 删除请求，包含文件/文件夹 ID
    """
    import shutil

    try:
        if not request.file_id:
            return DeleteFileResponse(
                success=False,
                file_id=request.file_id,
                error="文件 ID 不能为空",
            )

        workspace_dir = deps.config.workspace_dir
        target_path = Path(workspace_dir) / request.file_id

        # 安全检查：确保路径在 workspace 目录内
        try:
            target_path.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            return DeleteFileResponse(
                success=False,
                file_id=request.file_id,
                error="禁止访问工作区外的文件",
            )

        # 检查文件/文件夹是否存在
        if not target_path.exists():
            return DeleteFileResponse(
                success=False,
                file_id=request.file_id,
                error=f"文件或文件夹不存在: {request.file_id}",
            )

        # 删除文件或文件夹
        if target_path.is_file():
            target_path.unlink()
            logger.info(f"文件删除成功: {request.file_id}")
        else:
            # 递归删除文件夹
            shutil.rmtree(target_path)
            logger.info(f"文件夹删除成功: {request.file_id}")

        return DeleteFileResponse(
            success=True,
            file_id=request.file_id,
        )

    except Exception as e:
        logger.error(f"删除失败: {e}")
        return DeleteFileResponse(
            success=False,
            file_id=request.file_id,
            error=str(e),
        )


@router.post("/ai/process", response_model=DocumentAIResponse)
async def process_document_ai(request: DocumentAIRequest):
    """处理文档 AI 操作（非流式）

    Args:
        request: 文档 AI 操作请求
    """
    try:
        # 检查 LLM 是否可用
        if not deps.rag_service.llm_client.is_available():
            return DocumentAIResponse(
                success=False,
                response="",
                action=request.action.value,
                document_name=request.document_name,
                error="LLM服务不可用，请先配置 LLM",
            )

        # 获取对应的系统提示词
        system_prompt = _get_system_prompt(request.action, request.document_name)

        # 构建用户消息
        user_message = _build_user_message(request)

        # 调用 LLM
        response = deps.rag_service.llm_client.client.chat.completions.create(
            model=deps.rag_service.llm_client.model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.7,
        )

        # 记录 token 使用量
        if hasattr(response, "usage") and response.usage:
            log_token_usage(
                model=deps.rag_service.llm_client.model,
                input_tokens=response.usage.prompt_tokens,
                output_tokens=response.usage.completion_tokens,
                endpoint="workspace_ai_process",
                user_query=request.action.value,
                response_type="document_ai",
                feature_type="workspace_assistant",
                additional_info={
                    "action": request.action.value,
                    "document_name": request.document_name,
                    "document_length": len(request.document_content),
                },
            )

        result = response.choices[0].message.content.strip()

        return DocumentAIResponse(
            success=True,
            response=result,
            action=request.action.value,
            document_name=request.document_name,
        )

    except Exception as e:
        logger.error(f"文档 AI 处理失败: {e}")
        return DocumentAIResponse(
            success=False,
            response="",
            action=request.action.value,
            document_name=request.document_name,
            error=str(e),
        )


def _create_token_generator(request: DocumentAIRequest, messages: list[dict]):
    """创建流式 token 生成器"""

    def token_generator():
        try:
            response = deps.rag_service.llm_client.client.chat.completions.create(
                model=deps.rag_service.llm_client.model,
                messages=messages,
                temperature=0.7,
                stream=True,
                stream_options={"include_usage": True},
            )

            total_content = ""
            usage_info = None

            for chunk in response:
                if hasattr(chunk, "usage") and chunk.usage:
                    usage_info = chunk.usage

                if chunk.choices and len(chunk.choices) > 0 and chunk.choices[0].delta.content:
                    content = chunk.choices[0].delta.content
                    total_content += content
                    yield content

            _log_stream_usage(usage_info, request, total_content)

        except Exception as e:
            logger.error(f"流式生成失败: {e}")
            yield f"\n[错误] 流式生成失败: {e}"

    return token_generator


def _log_stream_usage(usage_info, request: DocumentAIRequest, total_content: str):
    """记录流式响应的 token 使用量"""
    if not usage_info:
        return

    try:
        log_token_usage(
            model=deps.rag_service.llm_client.model,
            input_tokens=usage_info.prompt_tokens,
            output_tokens=usage_info.completion_tokens,
            endpoint="workspace_ai_stream",
            user_query=request.action.value,
            response_type="stream",
            feature_type="workspace_assistant",
            additional_info={
                "action": request.action.value,
                "document_name": request.document_name,
                "document_length": len(request.document_content),
                "response_length": len(total_content),
            },
        )
    except Exception as log_error:
        logger.error(f"记录 token 使用量失败: {log_error}")


@router.post("/ai/stream")
async def process_document_ai_stream(request: DocumentAIRequest):
    """处理文档 AI 操作（流式输出）

    Args:
        request: 文档 AI 操作请求
    """
    try:
        if not deps.rag_service.llm_client.is_available():

            async def error_generator():
                yield "LLM服务不可用，请先配置 LLM"

            return StreamingResponse(error_generator(), media_type="text/plain; charset=utf-8")

        system_prompt = _get_system_prompt(request.action, request.document_name)
        user_message = _build_user_message(request)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ]

        token_generator = _create_token_generator(request, messages)
        headers = {"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}

        return StreamingResponse(
            token_generator(), media_type="text/plain; charset=utf-8", headers=headers
        )

    except Exception as e:
        logger.error(f"文档 AI 流式处理失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


def _get_system_prompt(action: DocumentAction, document_name: str | None = None) -> str:
    """获取对应操作的系统提示词

    Args:
        action: 操作类型
        document_name: 文档名称

    Returns:
        系统提示词
    """
    # 使用字典映射替代多个 if-elif 语句
    action_prompt_map = {
        DocumentAction.SUMMARIZE: "summarize",
        DocumentAction.IMPROVE: "improve",
        DocumentAction.EXPLAIN: "explain",
        DocumentAction.BEAUTIFY: "beautify",
        DocumentAction.EXPAND: "expand",
        DocumentAction.CONDENSE: "condense",
        DocumentAction.CORRECT: "correct",
        DocumentAction.TRANSLATE: "translate",
    }

    # 特殊处理 CUSTOM 操作
    if action == DocumentAction.CUSTOM:
        prompt = get_prompt("workspace_assistant", "custom_chat")
        return prompt.replace("{document_name}", document_name or "未命名文档")

    prompt_key = action_prompt_map.get(action, "summarize")
    return get_prompt("workspace_assistant", prompt_key)


def _build_user_message(request: DocumentAIRequest) -> str:
    """构建用户消息

    Args:
        request: 文档 AI 操作请求

    Returns:
        用户消息
    """
    # 文本编辑操作（美化、扩写、缩写、修正、翻译）只需要发送选中的文本
    text_edit_actions = [
        DocumentAction.BEAUTIFY,
        DocumentAction.EXPAND,
        DocumentAction.CONDENSE,
        DocumentAction.CORRECT,
        DocumentAction.TRANSLATE,
    ]

    if request.action in text_edit_actions:
        # 对于文本编辑操作，直接发送文本内容
        return request.document_content
    elif request.action == DocumentAction.CUSTOM and request.custom_prompt:
        # 自定义对话模式
        return f"""**文档内容：**
```
{request.document_content}
```

**用户问题：**
{request.custom_prompt}"""
    else:
        # 预定义操作模式（总结、改进、解释）
        doc_name_str = f"**文件名：** {request.document_name}\n\n" if request.document_name else ""
        return f"""{doc_name_str}**文档内容：**
```
{request.document_content}
```"""
